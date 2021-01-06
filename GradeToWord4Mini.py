#读取docx中的文本代码示例
import docx

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import Inches
from docx.shared import Pt

file=docx.Document(r"C:\Users\xulia\Desktop\1712041034-柯志宣-18软件工程（嵌入式培养）（1）.docx")
print("段落数:"+str(len(file.paragraphs))) #输出段落数

print(len("实验项目名称： Linux和Hadoop基础实验	"))

#输出每一段的内容
for para in file.paragraphs:
    if '实验成绩：' in para.text and '实验日期：' in para.text:
        print(para.text)
        text=para.text.rstrip()
        para.text=""
        index = text.index('实验成绩：')
        date_index=text.index('实验日期：')+5
        date=text[date_index:index].strip()+"  "

        run1=para.add_run('实验日期：')
        run1.font.name = '宋体'
        run1.font.underline = False  # 下划线
        run1.font.size = Pt(14)  # 字体大小

        run2 = para.add_run(date)
        run2.font.name = '宋体'
        run2.font.underline = True  # 下划线
        run2.font.size = Pt(14)  # 字体大小

        run2 = para.add_run(" "*(35-len(date)-5))
        run2.font.name = '宋体'
        run2.font.underline = False  # 下划线
        run2.font.size = Pt(14)  # 字体大小

        run3 = para.add_run('实验成绩：')
        run3.font.name = '宋体'
        run3.font.underline = False  # 下划线
        run3.font.size = Pt(14)  # 字体大小

        run4 = para.add_run("    "+str(80)+"    ")
        run4.font.name = '宋体'
        run4.font.underline = True  # 下划线
        run4.font.size = Pt(14)  # 字体大小
        # for run in para.runs:
        #     # 对于中文字体的设置
        #     run.font.name = '宋体'
        #     run.font.underline = False  # 下划线
        #     r = run._element.rPr.rFonts
        #     if str(80) in run.text or date in run.text:
        #         run.font.underline = True  # 下划线
        #     run.font.size = Pt(14)  # 字体大小
tables=file.tables
print("表格数:"+str(len(tables)))

grade_table=tables[0]

trun1=grade_table.rows[1].cells[5].paragraphs[0].add_run(str(80))
trun1.font.name = '宋体'
trun1.font.size = Pt(12)  # 字体大小


trun1=grade_table.rows[2].cells[5].paragraphs[0].add_run(str(80))
trun1.font.name = '宋体'
trun1.font.size = Pt(12)  # 字体大小

trun1=grade_table.rows[3].cells[5].paragraphs[0].add_run(str(80))
trun1.font.name = '宋体'
trun1.font.size = Pt(12)  # 字体大小

trun1=grade_table.rows[4].cells[5].paragraphs[0].add_run(str(80))
trun1.font.name = '宋体'
trun1.font.size = Pt(12)  # 字体大小

trun1=grade_table.rows[5].cells[6].paragraphs[0].add_run(str(80))
trun1.font.name = '宋体'
trun1.font.size = Pt(12)  # 字体大小
trun1.alignment = WD_ALIGN_PARAGRAPH.CENTER


summary_table=tables[1]
summary_table.rows[0].cells[0].paragraphs[0].text=""
srun=summary_table.rows[0].cells[0].paragraphs[0].add_run("该生能够较好地完成够较好地完成够较好地完成够较好地完成够较好地完成够较好地完成够较好地完成够较好地完成实验")
srun.font.name = '宋体'
srun.font.size = Pt(12)  # 字体大小
srun.first_line_indent =406400
summary_table.rows[0].cells[0].paragraphs[0].paragraph_format.first_line_indent = Inches(.25)

file.save((r"C:\Users\xulia\Desktop\1712041034-柯志宣-18软件工程（嵌入式培养）（1）-2.docx"))