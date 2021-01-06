#读取docx中的文本代码示例
import docx
import re

#获取文档
file=docx.Document(r"C:\Users\xulia\Desktop\1712041034-柯志宣-18软件工程（嵌入式培养）（1）.docx")
print("段落数:"+str(len(file.paragraphs))) #输出段落数


#输出每一段的内容
for para in file.paragraphs:
    if '实验成绩：' in para.text:
        print(para.text)
        para.text=para.text+str(80)

tables=file.tables
print("表格数:"+str(len(tables)))

grade_table=tables[0]

grade_table.rows[1].cells[5].text=str(80)
grade_table.rows[2].cells[5].text=str(80)
grade_table.rows[3].cells[5].text=str(80)
grade_table.rows[4].cells[5].text=str(80)
grade_table.rows[5].cells[6].text=str(80)

summary_table=tables[1]
summary_table.rows[0].cells[0].text='该生完成了报告'

file.save((r"C:\Users\xulia\Desktop\1712041034-柯志宣-18软件工程（嵌入式培养）（1）-2.docx"))